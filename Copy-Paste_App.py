import pyperclip
from docx import Document
import time
import re
import tkinter as tk
from threading import Thread

# Yeni bir Word belgesi oluştur
doc = Document()
doc.add_heading("Otomatik Yapıştırma", level=1)

# Önceki kopyalanan metni saklamak için
previous_text = ""
is_running = False

# Kopyalama işlemini başlatan fonksiyon
def start_copying():
    global is_running, previous_text
    is_running = True
    previous_text = ""  # Önceki metni sıfırla
    copy_thread = Thread(target=copy_loop)
    copy_thread.start()

# Kopyalama işlemini durduran fonksiyon
def stop_copying():
    global is_running
    is_running = False

# Kopyalama döngüsü
def copy_loop():
    global previous_text
    while is_running:
        # Panodaki metni al
        text = pyperclip.paste()
        
        # Eğer metin değiştiyse (yeni bir metin kopyalandıysa)
        if text != previous_text:
            previous_text = text
            
            # Satır aralarındaki fazla boşlukları kaldır ve kompakt hale getir
            formatted_text = re.sub(r'\n\s*', '\n', text).strip()
            
            # Yeni metni Word belgesine ekle
            doc.add_paragraph(formatted_text)
            
            # İstediğiniz ayrıcı metni ekleyin
            doc.add_paragraph("\n***\n")
            print("Yeni metin eklendi:", formatted_text)
            
            # Belgeyi kaydet
            doc.save("otomatik_yapistirma.docx")
        
        # Her 1 saniyede bir kontrol eder
        time.sleep(1)

# Tkinter arayüzü
window = tk.Tk()
window.title("Otomatik Kopyalama ve Yapıştırma")

# Başlat butonu
start_button = tk.Button(window, text="Başlat", command=start_copying)
start_button.pack(pady=10)

# Durdur butonu
stop_button = tk.Button(window, text="Durdur", command=stop_copying)
stop_button.pack(pady=10)

# Bilgilendirme etiketi
info_label = tk.Label(window, text="Metin kopyalamaya başlamak için 'Başlat'a tıklayın.")
info_label.pack(pady=20)

# Tkinter ana döngüsü
window.mainloop()
