<<<<<<< HEAD
import pyperclip
from docx import Document
import time
import re

# Yeni bir Word belgesi oluştur
doc = Document()
doc.add_heading("Otomatik Yapıştırma", level=1)

# Önceki kopyalanan metni saklamak
previous_text = ""

print("Otomatik kopyalama ve yapıştırma başlatıldı. Arka planda çalışıyor...")

while True:
    # Panodaki metni al
    text = pyperclip.paste()
    
    # Eğer metin değiştiyse (yeni bir metin kopyalandıysa)
    if text != previous_text:
        previous_text = text
        
        # Satır aralarındaki fazla boşlukları kaldır ve kompakt hale getir
        # Tüm satır başı boşlukları kaldırarak tek bir paragraf yapısında düzenler
        formatted_text = re.sub(r'\n\s*', '\n', text).strip()
        
        # Yeni metni Word belgesine ekle
        doc.add_paragraph(formatted_text)
        
        # İstediğiniz ayrıcı metni ekleyin
        doc.add_paragraph("\n\n***\n\n")
        print("Yeni metin eklendi:", formatted_text)
        
        # Belgeyi kaydet
        doc.save("otomatik_yapistirma.docx")
    
    # Her 1 saniyede bir kontrol eder
    time.sleep(1)
=======
import pyperclip
from docx import Document
import time
import re

# Yeni bir Word belgesi oluştur
doc = Document()
doc.add_heading("Otomatik Yapıştırma", level=1)

# Önceki kopyalanan metni saklamak
previous_text = ""

print("Otomatik kopyalama ve yapıştırma başlatıldı. Arka planda çalışıyor...")

while True:
    # Panodaki metni al
    text = pyperclip.paste()
    
    # Eğer metin değiştiyse (yeni bir metin kopyalandıysa)
    if text != previous_text:
        previous_text = text
        
        # Satır aralarındaki fazla boşlukları kaldır ve kompakt hale getir
        # Tüm satır başı boşlukları kaldırarak tek bir paragraf yapısında düzenler
        formatted_text = re.sub(r'\n\s*', '\n', text).strip()
        
        # Yeni metni Word belgesine ekle
        doc.add_paragraph(formatted_text)
        
        # İstediğiniz ayrıcı metni ekleyin
        doc.add_paragraph("\n\n***\n\n")
        print("Yeni metin eklendi:", formatted_text)
        
        # Belgeyi kaydet
        doc.save("otomatik_yapistirma.docx")
    
    # Her 1 saniyede bir kontrol eder
    time.sleep(1)
>>>>>>> 068d2118efd387fd41b94afd98d869e117a3337b
